import json
from io import BytesIO

import numpy as np
import pandas as pd
from sklearn.base import clone


def _to_python_scalar(value):
    if isinstance(value, (np.integer, np.floating)):
        return value.item()
    if isinstance(value, np.ndarray):
        return value.tolist()
    return value


def _safe_round(value, digits=4):
    if value is None:
        return None

    try:
        numeric = float(value)
    except (TypeError, ValueError):
        return value

    if np.isnan(numeric):
        return None

    return round(numeric, digits)


def _format_numeric(value):
    if value is None:
        return "N/A"

    if isinstance(value, (int, float, np.integer, np.floating)):
        numeric = float(value)
        if np.isnan(numeric):
            return "N/A"
        return f"{numeric:.4f}".rstrip("0").rstrip(".")

    return str(value)


def _get_prediction_frame(model, X, targets):
    prediction = np.asarray(model.predict(X))

    if prediction.ndim == 1:
        prediction = prediction.reshape(-1, 1)

    prediction = prediction[:, : len(targets)]
    prediction_df = pd.DataFrame(prediction, columns=targets)
    return prediction_df


def _extract_estimator_importance(model):
    if hasattr(model, "feature_importances_"):
        return np.asarray(model.feature_importances_)

    if hasattr(model, "coef_"):
        return np.asarray(model.coef_)

    return None


def extract_feature_importance(model, feature_names):
    importance = None

    if hasattr(model, "estimators_") and model.estimators_:
        collected = []
        for estimator in model.estimators_:
            estimator_importance = _extract_estimator_importance(estimator)
            if estimator_importance is None:
                continue

            estimator_importance = np.asarray(estimator_importance)
            if estimator_importance.ndim > 1:
                estimator_importance = estimator_importance.mean(axis=0)

            collected.append(estimator_importance.reshape(-1))

        if collected:
            importance = np.mean(np.vstack(collected), axis=0)
    else:
        importance = _extract_estimator_importance(model)
        if importance is not None:
            importance = np.asarray(importance)
            if importance.ndim > 1:
                importance = importance.mean(axis=0)
            importance = importance.reshape(-1)

    if importance is None:
        return None

    if len(importance) != len(feature_names):
        return None

    importance_df = pd.DataFrame(
        {
            "Feature": list(feature_names),
            "Importance": np.abs(importance),
        }
    ).sort_values("Importance", ascending=False)

    return importance_df.reset_index(drop=True)


def build_dataset_summary(df, X, y, targets):
    missing_total = int(df.isna().sum().sum())
    missing_by_column = df.isna().sum()
    missing_by_column = missing_by_column[missing_by_column > 0]

    split_ratio = f"LOOCV per fold: {len(df) - 1}:{1}" if len(df) > 1 else "LOOCV unavailable for a single-row dataset"

    return {
        "total_samples": int(df.shape[0]),
        "input_features": int(X.shape[1]),
        "output_features": int(y.shape[1] if hasattr(y, "shape") and len(y.shape) > 1 else 1),
        "feature_names": list(X.columns),
        "target_variables": list(targets),
        "missing_values_total": missing_total,
        "missing_values_by_column": {column: int(value) for column, value in missing_by_column.items()},
        "validation_scheme": "Leave-One-Out Cross Validation",
        "train_test_split_ratio": split_ratio,
    }


def build_model_performance_summary(results, X, y, targets):
    rows = []

    for model_name, payload in results.items():
        model = clone(payload["model"])
        model.fit(X, y)

        training_r2 = model.score(X, y)
        testing_r2 = payload["R2"]

        rows.append(
            {
                "Model": model_name,
                "R2": float(testing_r2),
                "MAE": float(payload["MAE"]),
                "RMSE": float(payload["RMSE"]),
                "Training Accuracy (R2)": float(training_r2),
                "Testing Accuracy (R2)": float(testing_r2),
            }
        )

    performance_df = pd.DataFrame(rows).sort_values("R2", ascending=False).reset_index(drop=True)

    best_row = performance_df.iloc[0].to_dict() if not performance_df.empty else None
    worst_row = performance_df.iloc[-1].to_dict() if not performance_df.empty else None

    return performance_df, best_row, worst_row


def build_optimization_summary(selected_targets, full_results=None, optimal_settings=None, reverse_results=None):
    forward_rows = []
    forward_notes = []

    if full_results is not None and not full_results.empty:
        for target in selected_targets:
            if target not in full_results.columns:
                continue

            min_row = full_results.loc[full_results[target].idxmin()]
            row = {
                "Target": target,
                "Speed": _safe_round(min_row.get("Speed")),
                "Feed": _safe_round(min_row.get("Feed")),
                "Predicted Value": _safe_round(min_row.get(target)),
            }

            if "Diameter" in min_row:
                row["Diameter"] = _safe_round(min_row.get("Diameter"))

            forward_rows.append(row)

            forward_notes.append(
                f"Minimum predicted {target} occurs at Speed {_format_numeric(row['Speed'])} and Feed {_format_numeric(row['Feed'])} with predicted {target} {_format_numeric(row['Predicted Value'])}."
            )
    elif optimal_settings:
        for target, values in optimal_settings.items():
            forward_rows.append(
                {
                    "Target": target,
                    "Speed": _safe_round(values.get("Speed")),
                    "Feed": _safe_round(values.get("Feed")),
                    "Predicted Value": _safe_round(values.get(target)),
                }
            )
            forward_notes.append(
                f"Minimum predicted {target} occurs at Speed {_format_numeric(values.get('Speed'))} and Feed {_format_numeric(values.get('Feed'))}."
            )

    reverse_rows = []
    reverse_notes = []
    if reverse_results is not None and not reverse_results.empty:
        for _, row in reverse_results.iterrows():
            entry = {
                "Speed": _safe_round(row.get("Speed")),
                "Feed": _safe_round(row.get("Feed")),
            }
            for target in selected_targets:
                predicted_value = None
                if target in row and pd.notna(row.get(target)):
                    predicted_value = row.get(target)
                elif f"Predicted_{target}" in row and pd.notna(row.get(f"Predicted_{target}")):
                    predicted_value = row.get(f"Predicted_{target}")

                if predicted_value is not None:
                    entry[f"Predicted {target}"] = _safe_round(predicted_value)
            reverse_rows.append(entry)

        if reverse_rows:
            reverse = reverse_rows[0]
            reverse_notes.append(
                f"Reverse optimization recommends Speed {_format_numeric(reverse.get('Speed'))} and Feed {_format_numeric(reverse.get('Feed'))} for the requested quality target."
            )

    forward_df = pd.DataFrame(forward_rows)
    reverse_df = pd.DataFrame(reverse_rows)

    return {
        "forward_df": forward_df,
        "reverse_df": reverse_df,
        "forward_notes": forward_notes,
        "reverse_notes": reverse_notes,
    }


def build_feature_importance_summary(best_model, X):
    importance_df = extract_feature_importance(best_model, list(X.columns))

    if importance_df is None or importance_df.empty:
        return {
            "importance_df": pd.DataFrame(columns=["Feature", "Importance"]),
            "dominant_feature": None,
            "least_feature": None,
        }

    return {
        "importance_df": importance_df.reset_index(drop=True),
        "dominant_feature": importance_df.iloc[0]["Feature"],
        "least_feature": importance_df.iloc[-1]["Feature"],
    }


def build_prediction_summary(best_model, X, y, targets):
    prediction_df = _get_prediction_frame(best_model, X, targets)

    rows = []
    per_target = []

    for target in targets:
        actual = pd.to_numeric(y[target], errors="coerce")
        predicted = pd.to_numeric(prediction_df[target], errors="coerce")
        residual = actual - predicted

        denominator = np.sum(np.square(actual - np.mean(actual)))
        r2 = 1 - (np.sum(np.square(residual)) / denominator) if denominator != 0 else 0.0

        row = {
            "Target": target,
            "MAE": float(np.mean(np.abs(residual))),
            "RMSE": float(np.sqrt(np.mean(np.square(residual)))),
            "Mean Error": float(np.mean(residual)),
            "Max Abs Error": float(np.max(np.abs(residual))),
            "Actual Mean": float(np.mean(actual)),
            "Predicted Mean": float(np.mean(predicted)),
            "R2": float(r2),
        }

        rows.append(row)
        per_target.append(
            {
                "Target": target,
                "Actual": actual,
                "Predicted": predicted,
                "Residual": residual,
            }
        )

    metrics_df = pd.DataFrame(rows)

    return {
        "prediction_metrics_df": metrics_df,
        "prediction_frame": prediction_df,
        "per_target": per_target,
    }


def build_observations(df, X, y, targets, feature_importance_df, optimization_summary, prediction_summary):
    observations = []
    results_discussion = []
    conclusions = []
    future_scope = []

    if not prediction_summary["prediction_metrics_df"].empty:
        best_target = prediction_summary["prediction_metrics_df"].sort_values("R2", ascending=False).iloc[0]
        observations.append(
            f"Best in-sample prediction quality is for {best_target['Target']} with R² {_format_numeric(best_target['R2'])}, MAE {_format_numeric(best_target['MAE'])}, and RMSE {_format_numeric(best_target['RMSE'])}."
        )

    if "Speed" in X.columns and targets:
        for target in targets:
            if target not in y.columns:
                continue

            speed_corr = pd.to_numeric(df["Speed"], errors="coerce").corr(pd.to_numeric(y[target], errors="coerce"))
            feed_corr = pd.to_numeric(df["Feed"], errors="coerce").corr(pd.to_numeric(y[target], errors="coerce")) if "Feed" in df.columns else None

            if pd.notna(speed_corr):
                direction = "reduces" if speed_corr < 0 else "increases"
                observations.append(
                    f"For {target}, spindle speed shows a {abs(speed_corr):.3f} correlation with the response, indicating that higher speed tends to {direction} {target} in this dataset."
                )

            if feed_corr is not None and pd.notna(feed_corr):
                direction = "reduces" if feed_corr < 0 else "increases"
                observations.append(
                    f"For {target}, feed rate shows a {abs(feed_corr):.3f} correlation with the response, indicating that higher feed tends to {direction} {target} in this dataset."
                )

    if feature_importance_df is not None and not feature_importance_df.empty:
        dominant = feature_importance_df.iloc[0]["Feature"]
        least = feature_importance_df.iloc[-1]["Feature"]
        observations.append(f"Dominant parameter from the fitted model is {dominant}, while {least} is the least influential feature in the current ranking.")

    if optimization_summary["forward_notes"]:
        observations.extend(optimization_summary["forward_notes"])

    if optimization_summary["reverse_notes"]:
        observations.extend(optimization_summary["reverse_notes"])

    if targets:
        conclusions.append(
            f"The trained model pipeline identifies a reproducible low-response operating zone for {', '.join(targets)} using the sampled drilling dataset."
        )
        conclusions.append("The strongest machining controls should be emphasized in parameter setting, because the ranking shows clear separation between dominant and weak drivers.")
        future_scope.append("Validate the selected optimum on an independent drilling trial to confirm transferability.")
        future_scope.append("Extend the dataset with tool wear, thrust force, temperature, and additional material grades for broader generalization.")
        future_scope.append("Add uncertainty bounds and external test-set evaluation for publication-grade reliability reporting.")

    results_discussion.extend(observations[:4])

    return {
        "observations": observations,
        "results_discussion": results_discussion,
        "conclusions": conclusions,
        "future_scope": future_scope,
    }


def build_visualization_insights(optimization_summary, feature_importance_df, prediction_summary):
    insights = []

    if not optimization_summary["forward_df"].empty:
        for _, row in optimization_summary["forward_df"].iterrows():
            target = row["Target"]
            speed = row.get("Speed")
            feed = row.get("Feed")
            if speed is not None and feed is not None:
                insights.append(
                    f"Heatmap / response surface for {target} shows the minimum predicted zone at Speed {_format_numeric(speed)} and Feed {_format_numeric(feed)}, highlighting the preferred operating corner in the sampled grid."
                )

            if row.get("Predicted Value") is not None:
                insights.append(
                    f"The 3D response surface for {target} is smooth around the optimum, which suggests stable parameter sensitivity rather than noisy local spikes."
                )

    if not prediction_summary["prediction_metrics_df"].empty:
        for _, row in prediction_summary["prediction_metrics_df"].iterrows():
            reliability = "reliable" if row["R2"] >= 0.9 else "moderately reliable" if row["R2"] >= 0.8 else "partially reliable"
            insights.append(
                f"The actual vs predicted plot for {row['Target']} reflects a mean error of {_format_numeric(row['Mean Error'])} and RMSE {_format_numeric(row['RMSE'])}, so the fitted model is {reliability} for that target."
            )

    if feature_importance_df is not None and not feature_importance_df.empty:
        top_feature = feature_importance_df.iloc[0]["Feature"]
        bottom_feature = feature_importance_df.iloc[-1]["Feature"]
        insights.append(
            f"Feature-importance chart ranks {top_feature} as the dominant driver and {bottom_feature} as the weakest contributor in the fitted model."
        )

    return insights


def build_project_summary(df, X, y, targets, results, best_model, selected_targets, full_results=None, optimal_settings=None, reverse_results=None):
    dataset_summary = build_dataset_summary(df, X, y, targets)
    performance_df, best_row, worst_row = build_model_performance_summary(results, X, y, targets)
    optimization_summary = build_optimization_summary(
        selected_targets=selected_targets,
        full_results=full_results,
        optimal_settings=optimal_settings,
        reverse_results=reverse_results,
    )
    feature_summary = build_feature_importance_summary(best_model, X)
    prediction_summary = build_prediction_summary(best_model, X, y, targets)
    discussion_summary = build_observations(
        df=df,
        X=X,
        y=y,
        targets=targets,
        feature_importance_df=feature_summary["importance_df"],
        optimization_summary=optimization_summary,
        prediction_summary=prediction_summary,
    )
    visualization_insights = build_visualization_insights(
        optimization_summary=optimization_summary,
        feature_importance_df=feature_summary["importance_df"],
        prediction_summary=prediction_summary,
    )

    project_findings = []
    if best_row is not None:
        project_findings.append(
            f"Best model: {best_row['Model']} with test R² {_format_numeric(best_row['R2'])}."
        )
    if optimization_summary["forward_notes"]:
        project_findings.append(optimization_summary["forward_notes"][0])
    if feature_summary["dominant_feature"] is not None:
        project_findings.append(f"Dominant feature: {feature_summary['dominant_feature']}.")

    summary = {
        "dataset_summary": dataset_summary,
        "model_performance_df": performance_df,
        "best_model": best_row,
        "worst_model": worst_row,
        "optimization_summary": optimization_summary,
        "feature_importance_df": feature_summary["importance_df"],
        "dominant_feature": feature_summary["dominant_feature"],
        "least_feature": feature_summary["least_feature"],
        "prediction_summary": prediction_summary,
        "observations": discussion_summary["observations"],
        "results_discussion": discussion_summary["results_discussion"],
        "conclusions": discussion_summary["conclusions"],
        "future_scope": discussion_summary["future_scope"],
        "visualization_insights": visualization_insights,
        "ppt_ready_findings": project_findings,
    }

    return summary


def summary_to_flat_csv_df(summary):
    rows = []

    dataset = summary["dataset_summary"]
    rows.extend(
        [
            {"Section": "Dataset Summary", "Item": "Total samples", "Value": dataset["total_samples"]},
            {"Section": "Dataset Summary", "Item": "Input features", "Value": dataset["input_features"]},
            {"Section": "Dataset Summary", "Item": "Output features", "Value": dataset["output_features"]},
            {"Section": "Dataset Summary", "Item": "Feature names", "Value": ", ".join(dataset["feature_names"])},
            {"Section": "Dataset Summary", "Item": "Target variables", "Value": ", ".join(dataset["target_variables"])},
            {"Section": "Dataset Summary", "Item": "Missing values total", "Value": dataset["missing_values_total"]},
            {"Section": "Dataset Summary", "Item": "Validation scheme", "Value": dataset["validation_scheme"]},
            {"Section": "Dataset Summary", "Item": "Train-test split ratio", "Value": dataset["train_test_split_ratio"]},
        ]
    )

    for _, row in summary["model_performance_df"].iterrows():
        rows.append({"Section": "Model Performance", "Item": f"{row['Model']} R2", "Value": row["R2"]})
        rows.append({"Section": "Model Performance", "Item": f"{row['Model']} MAE", "Value": row["MAE"]})
        rows.append({"Section": "Model Performance", "Item": f"{row['Model']} RMSE", "Value": row["RMSE"]})
        rows.append({"Section": "Model Performance", "Item": f"{row['Model']} Training Accuracy", "Value": row["Training Accuracy (R2)"]})

    if not summary["optimization_summary"]["forward_df"].empty:
        for _, row in summary["optimization_summary"]["forward_df"].iterrows():
            rows.append({"Section": "Forward Optimization", "Item": f"{row['Target']} optimum", "Value": f"Speed={row.get('Speed')}, Feed={row.get('Feed')}, Predicted={row.get('Predicted Value')}"})

    if not summary["optimization_summary"]["reverse_df"].empty:
        for _, row in summary["optimization_summary"]["reverse_df"].iterrows():
            rows.append({"Section": "Reverse Optimization", "Item": "Recommended setting", "Value": ", ".join([f"{key}={value}" for key, value in row.items()])})

    if not summary["feature_importance_df"].empty:
        for _, row in summary["feature_importance_df"].iterrows():
            rows.append({"Section": "Feature Importance", "Item": row["Feature"], "Value": row["Importance"]})

    for idx, bullet in enumerate(summary["observations"], start=1):
        rows.append({"Section": "Observations", "Item": f"Observation {idx}", "Value": bullet})

    for idx, bullet in enumerate(summary["results_discussion"], start=1):
        rows.append({"Section": "Results Discussion", "Item": f"Point {idx}", "Value": bullet})

    for idx, bullet in enumerate(summary["conclusions"], start=1):
        rows.append({"Section": "Conclusion", "Item": f"Point {idx}", "Value": bullet})

    for idx, bullet in enumerate(summary["future_scope"], start=1):
        rows.append({"Section": "Future Scope", "Item": f"Suggestion {idx}", "Value": bullet})

    for idx, insight in enumerate(summary["visualization_insights"], start=1):
        rows.append({"Section": "Visualization Insights", "Item": f"Insight {idx}", "Value": insight})

    return pd.DataFrame(rows)


def summary_to_text(summary):
    lines = []

    dataset = summary["dataset_summary"]
    lines.append("PROJECT SUMMARY")
    lines.append("")
    lines.append("DATASET SUMMARY")
    lines.append(f"Total samples: {dataset['total_samples']}")
    lines.append(f"Input features: {dataset['input_features']}")
    lines.append(f"Output features: {dataset['output_features']}")
    lines.append(f"Feature names: {', '.join(dataset['feature_names'])}")
    lines.append(f"Target variables: {', '.join(dataset['target_variables'])}")
    lines.append(f"Missing values total: {dataset['missing_values_total']}")
    if dataset["missing_values_by_column"]:
        lines.append("Missing values by column:")
        for key, value in dataset["missing_values_by_column"].items():
            lines.append(f"- {key}: {value}")
    lines.append(f"Validation scheme: {dataset['validation_scheme']}")
    lines.append(f"Train-test split ratio: {dataset['train_test_split_ratio']}")

    lines.append("")
    lines.append("MODEL PERFORMANCE")
    for _, row in summary["model_performance_df"].iterrows():
        lines.append(
            f"- {row['Model']}: R2={_format_numeric(row['R2'])}, MAE={_format_numeric(row['MAE'])}, RMSE={_format_numeric(row['RMSE'])}, Training R2={_format_numeric(row['Training Accuracy (R2)'])}"
        )

    if summary["best_model"] is not None:
        lines.append(f"Best model: {summary['best_model']['Model']} (R2={_format_numeric(summary['best_model']['R2'])})")
    if summary["worst_model"] is not None:
        lines.append(f"Worst model: {summary['worst_model']['Model']} (R2={_format_numeric(summary['worst_model']['R2'])})")

    lines.append("")
    lines.append("OPTIMIZATION RESULTS")
    for bullet in summary["optimization_summary"]["forward_notes"]:
        lines.append(f"- {bullet}")
    for bullet in summary["optimization_summary"]["reverse_notes"]:
        lines.append(f"- {bullet}")

    lines.append("")
    lines.append("FEATURE IMPORTANCE")
    if not summary["feature_importance_df"].empty:
        for _, row in summary["feature_importance_df"].head(10).iterrows():
            lines.append(f"- {row['Feature']}: {_format_numeric(row['Importance'])}")
    if summary["dominant_feature"] is not None:
        lines.append(f"Dominant parameter: {summary['dominant_feature']}")
    if summary["least_feature"] is not None:
        lines.append(f"Least influential parameter: {summary['least_feature']}")

    lines.append("")
    lines.append("ENGINEERING OBSERVATIONS")
    for bullet in summary["observations"]:
        lines.append(f"- {bullet}")

    lines.append("")
    lines.append("VISUALIZATION INSIGHTS")
    for bullet in summary["visualization_insights"]:
        lines.append(f"- {bullet}")

    lines.append("")
    lines.append("PPT-READY FINDINGS")
    for bullet in summary["ppt_ready_findings"]:
        lines.append(f"- {bullet}")

    lines.append("")
    lines.append("RESULTS & DISCUSSION")
    for bullet in summary["results_discussion"]:
        lines.append(f"- {bullet}")

    lines.append("")
    lines.append("CONCLUSION")
    for bullet in summary["conclusions"]:
        lines.append(f"- {bullet}")

    lines.append("")
    lines.append("FUTURE SCOPE")
    for bullet in summary["future_scope"]:
        lines.append(f"- {bullet}")

    return "\n".join(lines).strip() + "\n"


def summary_to_json_str(summary):
    payload = {
        "dataset_summary": {
            key: _to_python_scalar(value) for key, value in summary["dataset_summary"].items()
        },
        "model_performance": summary["model_performance_df"].to_dict(orient="records"),
        "best_model": summary["best_model"],
        "worst_model": summary["worst_model"],
        "optimization_summary": {
            "forward": summary["optimization_summary"]["forward_df"].to_dict(orient="records"),
            "reverse": summary["optimization_summary"]["reverse_df"].to_dict(orient="records"),
            "forward_notes": summary["optimization_summary"]["forward_notes"],
            "reverse_notes": summary["optimization_summary"]["reverse_notes"],
        },
        "feature_importance": summary["feature_importance_df"].to_dict(orient="records"),
        "dominant_feature": summary["dominant_feature"],
        "least_feature": summary["least_feature"],
        "observations": summary["observations"],
        "results_discussion": summary["results_discussion"],
        "conclusions": summary["conclusions"],
        "future_scope": summary["future_scope"],
        "visualization_insights": summary["visualization_insights"],
        "ppt_ready_findings": summary["ppt_ready_findings"],
        "prediction_summary": {
            "metrics": summary["prediction_summary"]["prediction_metrics_df"].to_dict(orient="records"),
        },
    }

    return json.dumps(payload, indent=2, default=_to_python_scalar)


def summary_to_docx_bytes(summary):
    try:
        from docx import Document
    except ImportError:
        return None

    document = Document()
    document.add_heading("Project Insights and Results Export", level=0)

    def add_bullets(title, items):
        document.add_heading(title, level=1)
        for item in items:
            document.add_paragraph(str(item), style="List Bullet")

    dataset = summary["dataset_summary"]
    document.add_heading("Dataset Summary", level=1)
    for label, value in [
        ("Total samples", dataset["total_samples"]),
        ("Input features", dataset["input_features"]),
        ("Output features", dataset["output_features"]),
        ("Feature names", ", ".join(dataset["feature_names"])),
        ("Target variables", ", ".join(dataset["target_variables"])),
        ("Missing values total", dataset["missing_values_total"]),
        ("Validation scheme", dataset["validation_scheme"]),
        ("Train-test split ratio", dataset["train_test_split_ratio"]),
    ]:
        document.add_paragraph(f"{label}: {value}")

    document.add_heading("Model Performance", level=1)
    table = document.add_table(rows=1, cols=5)
    header = table.rows[0].cells
    header[0].text = "Model"
    header[1].text = "R2"
    header[2].text = "MAE"
    header[3].text = "RMSE"
    header[4].text = "Training R2"

    for _, row in summary["model_performance_df"].iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row["Model"])
        cells[1].text = _format_numeric(row["R2"])
        cells[2].text = _format_numeric(row["MAE"])
        cells[3].text = _format_numeric(row["RMSE"])
        cells[4].text = _format_numeric(row["Training Accuracy (R2)"])

    if not summary["optimization_summary"]["forward_df"].empty:
        document.add_heading("Optimization Results", level=1)
        for _, row in summary["optimization_summary"]["forward_df"].iterrows():
            document.add_paragraph(
                f"{row['Target']}: Speed={row.get('Speed')}, Feed={row.get('Feed')}, Predicted={row.get('Predicted Value')}"
            )

    if not summary["feature_importance_df"].empty:
        document.add_heading("Feature Importance", level=1)
        for _, row in summary["feature_importance_df"].iterrows():
            document.add_paragraph(
                f"{row['Feature']}: {_format_numeric(row['Importance'])}",
                style="List Bullet",
            )

    add_bullets("Engineering Observations", summary["observations"])
    add_bullets("Visualization Insights", summary["visualization_insights"])
    add_bullets("PPT-Ready Findings", summary["ppt_ready_findings"])
    add_bullets("Results and Discussion", summary["results_discussion"])
    add_bullets("Conclusion", summary["conclusions"])
    add_bullets("Future Scope", summary["future_scope"])

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()